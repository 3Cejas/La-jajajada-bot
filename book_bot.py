import tweepy, re, time
from access import *
from random import randrange
import docx
import random
import os
import glob
# Setup API:
def twitter_setup():
    # Authenticate and access using keys:
    auth = tweepy.OAuthHandler(CONSUMER_KEY, CONSUMER_SECRET)
    auth.set_access_token(ACCESS_TOKEN, ACCESS_SECRET)
    # Return API access:
    api=tweepy.API(auth)
    return api

#Funcion que elege una carpeta random
def choosefolder(path):
    raiz=os.listdir(path)  
    #Esto es si hay carpetas que no te interesan
    #raiz = glob.glob('cartas')+glob.glob('poemas')
    raiz = random.choice(raiz)
    return raiz

#Funcion que elige un archivo de texto de manera random
def choosetext(path):
    dire = random.choice(os.listdir(path))
    return dire

# Funcion que lee el archivo word a txt
def readtxt():

    doc = docx.Document("La jajajada bot.docx")
    #doc = docx.Document("escritos/poemas/nuevo.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
#Funcion que encuentra mayusculas
def mayuscula(letra):
    n = letra
    if n == n.lower():
        if n in ("a","e","i","o","u"):
            res= True;
            
        else:
            res= True
            
    else:
        if n in ("A","E","I","O","U"):
            
            res= False
        else:
            res= False
            
        
    return res
# Function to extract status (will return status for post):
def extract_status(path=None):
    # No path => return "No book opened!"
    fuente="escritos/"
    carpeta=choosefolder(fuente)
    print(carpeta)
    archivo=choosetext(fuente+carpeta)
    print(archivo)
    path=fuente+carpeta+"/"+archivo
    print (path)
    if not path:
        return "No se abrio el libro"

    # Try to search a sentence in book:
    try:
        # Open and read textbook:
       print (readtxt(path))
       text=readtxt(path)
       return search_sentence(text)
        # If successfuly read, search sentence:
        

    except:
        # Book not found:
            return "Libro no encontrado"

# Function to search a sentence in book:
def search_sentence(text):
    #Inicializa status:
    index = randrange(0,len(text))
    status = 400
    encontrado = False
    while not (5 < status <= 278):
        #bucle para encontrar el inicio con mayuscula
        while index >= 0 and not(encontrado):
            if(text[index] == "."):
                encontrado = True
                
            if(index==0):
                encontrado = True
            index-=1
        #generamos el final aleatorio
        index +=3
        indexfinal=index+ randrange(40,278)
        status = len(text[index:indexfinal])
    # Replace breaks w/spaces:
    sentence = text[index:indexfinal]
    
    return sentence
#Función que acorta el texto para que la última palabra no esté cortada
def acortar (tweet):
    indice = len(tweet)
    while indice > 0:

        if(tweet[indice-1] == "." or tweet[indice-1] == ";" or tweet[indice-1] == "?" or tweet[indice-1] == "!" or tweet[indice-1] == ")" or tweet[indice-1] == "]" or tweet[indice-1] == "»"):
            
            tweet = tweet[:indice]
            return tweet
        indice -= 1
    
    return "-1"
if __name__ == '__main__':
    # Setup Twitter API:
    bot = twitter_setup()

    # Set waiting time:
    segs = 86400

    # Eternal posting:
    while True:
        
        # Extract status:
        status = search_sentence(readtxt())
        status = acortar(status.strip())
        while status == "-1":
            status = search_sentence(readtxt())
            status = acortar(status.strip())
        print (status)
        # Try to post status:
        try:

            #bot.update_status(status.strip())
            print("Successfully posted.")

        except tweepy.TweepError as e:
            print(e.reason)
        # Wait till next sentence extraction:
        time.sleep(86400)